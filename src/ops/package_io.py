"""Filesystem and DOCX package I/O helpers."""

from __future__ import annotations

import datetime as dt
import os
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from config import BACKUP_DIRNAME, DEFAULT_DIR


def resolve_path(value: str) -> Path:
    """Resolve a user-provided path against the configured default directory.

    Args:
        value: Absolute or relative filesystem path.

    Returns:
        Absolute normalized path.
    """

    path = Path(value)
    if not path.is_absolute():
        path = DEFAULT_DIR / path
    return path.resolve()


def ensure_docx(path: Path) -> None:
    """Validate that a path points to a ``.docx`` file.

    Args:
        path: Filesystem path to validate.

    Raises:
        ValueError: The path does not use the ``.docx`` suffix.
    """

    if path.suffix.lower() != ".docx":
        raise ValueError(f"Only .docx files are supported: {path}")


def clean_text(value: str) -> str:
    """Normalize text extracted from DOCX/XML sources.

    Args:
        value: Raw extracted text.

    Returns:
        Cleaned text with control markers removed.
    """

    return value.replace("\r", "").replace("\x07", "").replace("\x05", "").strip()


def load_document(path_value: str) -> tuple[Any, Path]:
    """Load an existing DOCX file through ``python-docx``.

    Args:
        path_value: Absolute or relative DOCX path.

    Returns:
        Loaded ``Document`` instance and resolved source path.

    Raises:
        FileNotFoundError: The DOCX file does not exist.
    """

    path = resolve_path(path_value)
    if not path.exists():
        raise FileNotFoundError(f"Document was not found: {path}")
    ensure_docx(path)
    return Document(str(path)), path


def load_or_create_document(path_value: str) -> tuple[Any, Path, bool]:
    """Load an existing DOCX file or create a new empty document.

    Args:
        path_value: Absolute or relative DOCX path.

    Returns:
        Tuple of document instance, resolved path and creation flag.
    """

    path = resolve_path(path_value)
    ensure_docx(path)
    if path.exists():
        return Document(str(path)), path, False
    return Document(), path, True


def backup_if_overwriting(source_path: Path, target: Path) -> None:
    """Create a timestamped backup when saving over the source DOCX file.

    Args:
        source_path: Original document path.
        target: Final save destination.
    """

    if target == source_path and source_path.exists():
        backup_dir = source_path.parent / BACKUP_DIRNAME
        backup_dir.mkdir(parents=True, exist_ok=True)
        timestamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
        backup_path = backup_dir / f"{source_path.stem}-{timestamp}{source_path.suffix}"
        shutil.copy2(source_path, backup_path)


def save_document(doc: Any, source_path: Path, output_path: str | None) -> Path:
    """Save a ``python-docx`` document with optional backup handling.

    Args:
        doc: ``python-docx`` document instance to save.
        source_path: Original path used to load the document.
        output_path: Optional alternate output path.

    Returns:
        Final saved path.
    """

    target = resolve_path(output_path) if output_path else source_path
    ensure_docx(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    backup_if_overwriting(source_path, target)
    doc.save(str(target))
    return target


def read_zip_xml(path: Path, part_name: str) -> etree._Element | None:
    """Read and parse an XML part from a DOCX zip package.

    Args:
        path: DOCX package path.
        part_name: Internal zip member path such as ``word/document.xml``.

    Returns:
        Parsed XML root element or ``None`` when the part is missing.
    """

    with zipfile.ZipFile(path, "r") as package:
        try:
            data = package.read(part_name)
        except KeyError:
            return None
    return etree.fromstring(data)


def serialize_xml(root: etree._Element) -> bytes:
    """Serialize an XML root node back to standalone UTF-8 bytes.

    Args:
        root: XML root element to serialize.

    Returns:
        Serialized XML bytes.
    """

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def save_zip_parts(source_path: Path, output_path: str | None, parts: dict[str, bytes]) -> Path:
    """Write selected DOCX package parts back into a new or existing zip archive.

    Args:
        source_path: Existing DOCX package to copy unchanged parts from.
        output_path: Optional alternate output path.
        parts: Mapping of zip member names to replacement bytes.

    Returns:
        Final saved DOCX path.
    """

    target = resolve_path(output_path) if output_path else source_path
    ensure_docx(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    backup_if_overwriting(source_path, target)

    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(target.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(source_path, "r") as src, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            copied: set[str] = set()
            for item in src.infolist():
                if item.filename in parts:
                    dst.writestr(item, parts[item.filename])
                else:
                    dst.writestr(item, src.read(item.filename))
                copied.add(item.filename)
            for part_name, data in parts.items():
                if part_name not in copied:
                    dst.writestr(part_name, data)
        shutil.move(str(temp_path), str(target))
    finally:
        if temp_path.exists():
            temp_path.unlink()
    return target


