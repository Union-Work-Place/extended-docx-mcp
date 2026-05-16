"""Document structure, formatting and serialization helpers for ``python-docx``."""

from __future__ import annotations

from typing import Any

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from config import (
    PAPER_SIZES_INCHES,
    PX_ALIGNMENTS,
    PX_ORIENTATIONS,
    PX_SECTION_STARTS,
    PX_STYLE_TYPES,
    PX_TABLE_ALIGNMENTS,
    PX_UNDERLINES,
    STRUCTURED_BLOCK_TYPES,
)
from ops.package_io import clean_text
from ops.text_ops import enum_name, normalize_mapping_value


def parent_element(parent: Any) -> Any:
    """Return the OOXML container element for a document-like object.

    Args:
        parent: Document, cell or paragraph parent object.

    Returns:
        Underlying lxml-backed container element.
    """

    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        return parent.element.body
    if hasattr(parent, "_tc"):
        return parent._tc
    return parent.element


def iter_paragraphs_in_parent(parent: Any, seen: set[int]) -> list[Paragraph]:
    """Recursively collect paragraphs from a document or table cell container.

    Args:
        parent: ``python-docx`` container object.
        seen: Set of OOXML node identities already visited.

    Returns:
        Paragraph objects in document order.
    """

    paragraphs: list[Paragraph] = []
    for child in parent_element(parent).iterchildren():
        if child.tag == qn("w:p"):
            key = id(child)
            if key not in seen:
                seen.add(key)
                paragraphs.append(Paragraph(child, parent))
        elif child.tag == qn("w:tbl"):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(iter_paragraphs_in_parent(cell, seen))
    return paragraphs


def iter_paragraphs(doc: Any) -> list[Paragraph]:
    """Return all document paragraphs, including those nested in tables.

    Args:
        doc: ``python-docx`` document instance.

    Returns:
        Paragraph list in document order.
    """

    return iter_paragraphs_in_parent(doc, set())


def iter_tables_in_parent(parent: Any, seen: set[int]) -> list[Table]:
    """Recursively collect tables from a document or nested cell container.

    Args:
        parent: ``python-docx`` container object.
        seen: Set of OOXML node identities already visited.

    Returns:
        Table objects in document order.
    """

    tables: list[Table] = []
    for child in parent_element(parent).iterchildren():
        if child.tag == qn("w:tbl"):
            key = id(child)
            if key not in seen:
                seen.add(key)
                table = Table(child, parent)
                tables.append(table)
                for row in table.rows:
                    for cell in row.cells:
                        tables.extend(iter_tables_in_parent(cell, seen))
    return tables


def iter_tables(doc: Any) -> list[Table]:
    """Return all document tables, including nested tables.

    Args:
        doc: ``python-docx`` document instance.

    Returns:
        Table list in document order.
    """

    return iter_tables_in_parent(doc, set())


def find_paragraph(doc: Any, paragraph_index: int | None, anchor_text: str | None) -> tuple[Paragraph, int]:
    """Locate a paragraph either by absolute index or by anchor text.

    Args:
        doc: ``python-docx`` document instance.
        paragraph_index: Optional absolute paragraph index.
        anchor_text: Optional text snippet to search for.

    Returns:
        Tuple of matching paragraph and its absolute index.
    """

    paragraphs = iter_paragraphs(doc)
    if paragraph_index is not None:
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            raise IndexError(f"Paragraph index out of range: {paragraph_index}")
        return paragraphs[paragraph_index], paragraph_index
    if anchor_text:
        for index, paragraph in enumerate(paragraphs):
            if anchor_text in paragraph.text:
                return paragraph, index
        raise ValueError(f"Anchor text was not found in document: {anchor_text}")
    raise ValueError("Provide paragraph_index or anchor_text")


def paragraph_to_dict(paragraph: Paragraph, index: int, include_runs: bool) -> dict[str, Any]:
    """Convert a paragraph into an MCP-friendly dictionary.

    Args:
        paragraph: Paragraph to serialize.
        index: Absolute paragraph index.
        include_runs: Whether to include run-level metadata.

    Returns:
        Serialized paragraph descriptor.
    """

    paragraph_format = paragraph.paragraph_format
    result: dict[str, Any] = {
        "index": index,
        "text": paragraph.text,
        "style_name": paragraph.style.name if paragraph.style else None,
        "alignment": enum_name(paragraph_format.alignment, PX_ALIGNMENTS),
        "keep_with_next": paragraph_format.keep_with_next,
        "left_indent_points": paragraph_format.left_indent.pt if paragraph_format.left_indent else None,
        "right_indent_points": paragraph_format.right_indent.pt if paragraph_format.right_indent else None,
        "first_line_indent_points": paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else None,
        "space_before_points": paragraph_format.space_before.pt if paragraph_format.space_before else None,
        "space_after_points": paragraph_format.space_after.pt if paragraph_format.space_after else None,
    }
    if include_runs:
        result["runs"] = [
            {
                "index": run_index,
                "text": run.text,
                "font_name": run.font.name,
                "font_size_points": run.font.size.pt if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "underline": enum_name(run.underline, PX_UNDERLINES),
                "all_caps": run.font.all_caps,
            }
            for run_index, run in enumerate(paragraph.runs)
        ]
    return result


def table_to_dict(table: Table, index: int, include_cells: bool) -> dict[str, Any]:
    """Convert a table into an MCP-friendly dictionary.

    Args:
        table: Table to serialize.
        index: Absolute table index.
        include_cells: Whether to include per-row cell values.

    Returns:
        Serialized table descriptor.
    """

    row_data: list[dict[str, Any]] = []
    if include_cells:
        for row_index, row in enumerate(table.rows):
            row_data.append({"row_index": row_index, "cells": [clean_text(cell.text) for cell in row.cells]})
    rows = len(table.rows)
    columns = max((len(row.cells) for row in table.rows), default=0)
    text = "\n".join(cell for row in row_data for cell in row.get("cells", []))
    return {
        "index": index,
        "rows": rows,
        "columns": columns,
        "style_name": table.style.name if table.style else None,
        "alignment": enum_name(table.alignment, PX_TABLE_ALIGNMENTS),
        "text": clean_text(text)[:500],
        "row_data": row_data,
    }


def apply_paragraph_format(paragraph: Paragraph, block: dict[str, Any], default_style_name: str | None = None) -> None:
    """Apply paragraph formatting options from a structured block payload.

    Args:
        paragraph: Paragraph to mutate.
        block: User-provided formatting block.
        default_style_name: Optional fallback style name.
    """

    style_name = block.get("style_name") or default_style_name
    if style_name:
        paragraph.style = str(style_name)
    paragraph_format = paragraph.paragraph_format
    if block.get("alignment") is not None:
        paragraph_format.alignment = normalize_mapping_value(str(block["alignment"]), PX_ALIGNMENTS, "alignment")
    if block.get("keep_with_next") is not None:
        paragraph_format.keep_with_next = bool(block["keep_with_next"])
    if block.get("first_line_indent_points") is not None:
        paragraph_format.first_line_indent = Pt(float(block["first_line_indent_points"]))
    if block.get("left_indent_points") is not None:
        paragraph_format.left_indent = Pt(float(block["left_indent_points"]))
    if block.get("right_indent_points") is not None:
        paragraph_format.right_indent = Pt(float(block["right_indent_points"]))
    if block.get("space_before_points") is not None:
        paragraph_format.space_before = Pt(float(block["space_before_points"]))
    if block.get("space_after_points") is not None:
        paragraph_format.space_after = Pt(float(block["space_after_points"]))


def apply_run_format(run: Any, run_data: dict[str, Any]) -> None:
    """Apply run formatting options from a structured run payload.

    Args:
        run: ``python-docx`` run object to mutate.
        run_data: User-provided formatting descriptor.
    """

    if run_data.get("font_name") is not None:
        run.font.name = str(run_data["font_name"])
    if run_data.get("font_size_points") is not None:
        run.font.size = Pt(float(run_data["font_size_points"]))
    if run_data.get("bold") is not None:
        run.bold = bool(run_data["bold"])
    if run_data.get("italic") is not None:
        run.italic = bool(run_data["italic"])
    if run_data.get("underline") is not None:
        run.underline = normalize_mapping_value(str(run_data["underline"]), PX_UNDERLINES, "underline")
    if run_data.get("all_caps") is not None:
        run.font.all_caps = bool(run_data["all_caps"])


def populate_paragraph(paragraph: Paragraph, block: dict[str, Any], default_style_name: str | None = None) -> Paragraph:
    """Populate an existing paragraph from a structured block payload.

    Args:
        paragraph: Paragraph to populate.
        block: Structured paragraph or heading block.
        default_style_name: Optional fallback style name.

    Returns:
        The populated paragraph object.
    """

    apply_paragraph_format(paragraph, block, default_style_name)
    runs = block.get("runs")
    if runs is not None:
        if not isinstance(runs, list):
            raise ValueError("Paragraph runs must be a list")
        for run_data in runs:
            if not isinstance(run_data, dict):
                raise ValueError("Each run must be an object")
            run = paragraph.add_run(str(run_data.get("text", "")))
            apply_run_format(run, run_data)
    else:
        paragraph.add_run(str(block.get("text", "")))
    return paragraph


def write_paragraph_block(doc: Any, block: dict[str, Any], default_style_name: str | None = None) -> Paragraph:
    """Append a structured paragraph block to a document.

    Args:
        doc: ``python-docx`` document instance.
        block: Structured paragraph or heading block.
        default_style_name: Optional fallback style name.

    Returns:
        Newly created paragraph object.
    """

    return populate_paragraph(doc.add_paragraph(), block, default_style_name)


def write_cell_value(cell: Any, cell_value: Any) -> None:
    """Replace the visible content of a table cell.

    Args:
        cell: ``python-docx`` cell object to mutate.
        cell_value: Raw text or structured run payload.
    """

    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    if isinstance(cell_value, dict):
        runs = cell_value.get("runs")
        if runs is not None:
            if not isinstance(runs, list):
                raise ValueError("Cell runs must be a list")
            for run_data in runs:
                if not isinstance(run_data, dict):
                    raise ValueError("Each cell run must be an object")
                run = paragraph.add_run(str(run_data.get("text", "")))
                apply_run_format(run, run_data)
            return
        paragraph.add_run(str(cell_value.get("text", "")))
        return
    paragraph.add_run(str(cell_value))


def write_table_block(doc: Any, block: dict[str, Any]) -> Table:
    """Append a structured table block to a document.

    Args:
        doc: ``python-docx`` document instance.
        block: Structured table descriptor.

    Returns:
        Newly created table object.
    """

    rows = block.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("Table block must contain a non-empty rows list")
    if not all(isinstance(row, list) and row for row in rows):
        raise ValueError("Each table row must be a non-empty list")
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            write_cell_value(table.cell(row_index, column_index), value)
    if block.get("style_name") is not None:
        table.style = str(block["style_name"])
    if block.get("alignment") is not None:
        table.alignment = normalize_mapping_value(str(block["alignment"]), PX_TABLE_ALIGNMENTS, "table alignment")
    return table


def insert_table_after(paragraph: Paragraph, data: list[list[str]]) -> Table:
    """Insert a table directly after a paragraph.

    Args:
        paragraph: Paragraph after which the table should be inserted.
        data: Two-dimensional table data matrix.

    Returns:
        Inserted table object.
    """

    if not data or not data[0]:
        raise ValueError("Table data must contain at least one row and one cell")
    document = paragraph.part.document
    column_count = max(len(row) for row in data)
    table = document.add_table(rows=len(data), cols=column_count)
    for row_index, row_values in enumerate(data):
        for cell_index in range(column_count):
            value = row_values[cell_index] if cell_index < len(row_values) else ""
            write_cell_value(table.cell(row_index, cell_index), value)
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Insert a paragraph directly after another paragraph.

    Args:
        paragraph: Paragraph after which to insert a new paragraph.
        text: Initial paragraph text.

    Returns:
        Newly inserted paragraph object.
    """

    inserted = OxmlElement("w:p")
    paragraph._p.addnext(inserted)
    new_paragraph = Paragraph(inserted, paragraph._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def normalize_paragraph_range(total: int, start_paragraph: int, end_paragraph: int) -> tuple[int, int]:
    """Validate and normalize an inclusive paragraph range."""

    if start_paragraph < 0 or end_paragraph < 0:
        raise IndexError("Paragraph range indices must be >= 0")
    if start_paragraph > end_paragraph:
        raise ValueError("start_paragraph must be <= end_paragraph")
    if end_paragraph >= total:
        raise IndexError(f"Paragraph index out of range: {end_paragraph}")
    return start_paragraph, end_paragraph


def insert_structured_block_after(paragraph: Paragraph, block: dict[str, Any], block_index: int) -> Any:
    """Insert one supported structured block directly after a paragraph.

    Args:
        paragraph: Paragraph after which the block should be inserted.
        block: Structured block descriptor.
        block_index: Block index used in validation messages.
    """

    block_type = str(block.get("type", "paragraph")).lower()
    if block_type == "paragraph":
        return populate_paragraph(insert_paragraph_after(paragraph, ""), block)
    if block_type == "heading":
        level = int(block.get("level", 1))
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        return populate_paragraph(insert_paragraph_after(paragraph, ""), block, default_style_name=f"Heading {level}")
    if block_type == "table":
        rows = block.get("rows")
        if not isinstance(rows, list) or not rows:
            raise ValueError("Table block must contain a non-empty rows list")
        inserted = insert_table_after(paragraph, rows)
        if block.get("style_name") is not None:
            inserted.style = str(block["style_name"])
        if block.get("alignment") is not None:
            inserted.alignment = normalize_mapping_value(str(block["alignment"]), PX_TABLE_ALIGNMENTS, "table alignment")
        return inserted
    raise ValueError(f"Unsupported block type at index {block_index}: {block_type}")


def validate_structured_blocks(blocks: list[dict[str, Any]]) -> None:
    """Validate the high-level list payload accepted by ``write_docx``.

    Args:
        blocks: Structured content blocks to validate.

    Raises:
        ValueError: One of the blocks has an invalid type or payload shape.
    """

    if not isinstance(blocks, list):
        raise ValueError("blocks must be a list")
    for index, block in enumerate(blocks):
        if not isinstance(block, dict):
            raise ValueError(f"Block at index {index} must be an object")
        block_type = str(block.get("type", "paragraph")).lower()
        if block_type not in STRUCTURED_BLOCK_TYPES:
            raise ValueError(f"Unsupported block type at index {index}: {block_type}")


def write_structured_block(doc: Any, block: dict[str, Any], index: int) -> None:
    """Write one structured block into a document.

    Args:
        doc: ``python-docx`` document instance.
        block: Structured block descriptor.
        index: Block index used in validation messages.
    """

    block_type = str(block.get("type", "paragraph")).lower()
    if block_type not in STRUCTURED_BLOCK_TYPES:
        raise ValueError(f"Unsupported block type at index {index}: {block_type}")
    if block_type == "paragraph":
        write_paragraph_block(doc, block)
        return
    if block_type == "heading":
        level = int(block.get("level", 1))
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        write_paragraph_block(doc, block, default_style_name=f"Heading {level}")
        return
    if block_type == "table":
        write_table_block(doc, block)
        return
    if block_type == "page_break":
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break()
        return
    break_name = str(block.get("break", "new_page"))
    doc.add_section(normalize_mapping_value(break_name, PX_SECTION_STARTS, "section break"))


def section_to_dict(section: Any, index: int) -> dict[str, Any]:
    """Serialize a document section into an MCP-friendly dictionary.

    Args:
        section: ``python-docx`` section object.
        index: Absolute section index.

    Returns:
        Serialized section descriptor.
    """

    return {
        "index": index,
        "orientation": enum_name(section.orientation, PX_ORIENTATIONS),
        "page_width_points": section.page_width.pt,
        "page_height_points": section.page_height.pt,
        "left_margin_points": section.left_margin.pt,
        "right_margin_points": section.right_margin.pt,
        "top_margin_points": section.top_margin.pt,
        "bottom_margin_points": section.bottom_margin.pt,
        "different_first_page_header_footer": section.different_first_page_header_footer,
        "section_start": enum_name(getattr(section, "start_type", None), PX_SECTION_STARTS),
    }


def update_section_page_setup(
    section: Any,
    orientation: str | None,
    paper_size: str | None,
    section_start: str | None,
    left_margin_points: float | None,
    right_margin_points: float | None,
    top_margin_points: float | None,
    bottom_margin_points: float | None,
    different_first_page_header_footer: bool | None,
) -> dict[str, bool]:
    """Apply editable page setup properties to a section.

    Args:
        section: ``python-docx`` section object.
        orientation: Optional public orientation value.
        paper_size: Optional public paper size value.
        section_start: Optional section start mode.
        left_margin_points: Optional left margin in points.
        right_margin_points: Optional right margin in points.
        top_margin_points: Optional top margin in points.
        bottom_margin_points: Optional bottom margin in points.
        different_first_page_header_footer: Optional header/footer flag.

    Returns:
        Mapping of ignored option flags for unsupported fields.
    """

    if orientation is not None:
        normalized_orientation = normalize_mapping_value(orientation, PX_ORIENTATIONS, "orientation")
        section.orientation = normalized_orientation
        width, height = section.page_width, section.page_height
        if normalized_orientation == PX_ORIENTATIONS["landscape"] and height > width:
            section.page_width, section.page_height = height, width
        if normalized_orientation == PX_ORIENTATIONS["portrait"] and width > height:
            section.page_width, section.page_height = height, width
    if paper_size is not None:
        normalized_paper = paper_size.lower()
        if normalized_paper not in PAPER_SIZES_INCHES:
            raise ValueError(f"Unsupported paper_size: {paper_size}")
        width, height = PAPER_SIZES_INCHES[normalized_paper]
        if section.orientation == PX_ORIENTATIONS["landscape"]:
            width, height = height, width
        section.page_width = Inches(width)
        section.page_height = Inches(height)
    if section_start is not None and hasattr(section, "start_type"):
        section.start_type = normalize_mapping_value(section_start, PX_SECTION_STARTS, "section_start")
    if left_margin_points is not None:
        section.left_margin = Pt(left_margin_points)
    if right_margin_points is not None:
        section.right_margin = Pt(right_margin_points)
    if top_margin_points is not None:
        section.top_margin = Pt(top_margin_points)
    if bottom_margin_points is not None:
        section.bottom_margin = Pt(bottom_margin_points)
    if different_first_page_header_footer is not None:
        section.different_first_page_header_footer = different_first_page_header_footer
    return {}


def style_to_dict(style: Any) -> dict[str, Any]:
    """Serialize a Word style object.

    Args:
        style: ``python-docx`` style object.

    Returns:
        Serialized style descriptor.
    """

    return {
        "name": style.name,
        "built_in": bool(style.builtin),
        "type": enum_name(style.type, PX_STYLE_TYPES),
        "base_style_name": style.base_style.name if style.base_style else None,
        "font_name": style.font.name,
        "font_size": style.font.size.pt if style.font.size else None,
        "bold": style.font.bold,
        "italic": style.font.italic,
    }


def ensure_style_type_is_paragraph(style: Any) -> None:
    """Validate that a style supports paragraph formatting settings.

    Args:
        style: ``python-docx`` style object to inspect.

    Raises:
        ValueError: The style is not a paragraph style.
    """

    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        raise ValueError("Paragraph formatting can only be applied to paragraph styles")
