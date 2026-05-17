"""Utilities for generating reproducible DOCX test fixtures."""

from __future__ import annotations

import asyncio
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from lxml import etree

from app import create_server
from config import W_NS


def _save(doc: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_simple(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("Alpha project budget is 100.")
    doc.add_paragraph("The draft conclusion needs review.")
    doc.add_paragraph("Final summary for regional team.")
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Q1"
    table.cell(0, 2).text = "Q2"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "120"
    table.cell(1, 2).text = "145"
    return _save(doc, path)


def build_with_styles(path: Path) -> Path:
    doc = Document()
    style = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(14)
    style.font.bold = True
    doc.add_heading("Styled document", level=1)
    paragraph = doc.add_paragraph("Styled callout paragraph.")
    paragraph.style = "Callout"
    doc.add_paragraph("Reference body paragraph.")
    return _save(doc, path)


def build_with_sections(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Section one starts here.")
    second = doc.add_section(WD_SECTION_START.NEW_PAGE)
    second.orientation = WD_ORIENT.LANDSCAPE
    doc.add_paragraph("Section two starts here.")
    return _save(doc, path)


def _rewrite_docx_xml(path: Path, part_name: str, mutator) -> Path:
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    Path(temp_name).unlink(missing_ok=True)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == part_name:
                    root = etree.fromstring(data)
                    mutator(root)
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                target.writestr(item, data)
        shutil.move(str(temp_path), str(path))
    finally:
        if temp_path.exists():
            temp_path.unlink()
    return path


def build_fractional_sections(path: Path) -> Path:
    build_with_sections(path)

    def _mutate(root: etree._Element) -> None:
        namespace = {"w": W_NS}
        section_nodes = root.xpath("//w:sectPr", namespaces=namespace)
        if not section_nodes:
            raise RuntimeError("Expected at least one section in test fixture")
        page_margins = section_nodes[-1].find("w:pgMar", namespaces=namespace)
        if page_margins is None:
            raise RuntimeError("Expected w:pgMar in section properties")
        page_margins.set(f"{{{W_NS}}}left", "1984.251968503937")

    return _rewrite_docx_xml(path, "word/document.xml", _mutate)


def build_fractional_paragraph_formatting(path: Path) -> Path:
    build_simple(path)

    def _mutate(root: etree._Element) -> None:
        namespace = {"w": W_NS}
        paragraph = root.xpath("//w:body/w:p[2]", namespaces=namespace)
        if not paragraph:
            raise RuntimeError("Expected a second body paragraph in test fixture")
        paragraph_properties = paragraph[0].find("w:pPr", namespaces=namespace)
        if paragraph_properties is None:
            paragraph_properties = etree.SubElement(paragraph[0], f"{{{W_NS}}}pPr")
        indent = paragraph_properties.find("w:ind", namespaces=namespace)
        if indent is None:
            indent = etree.SubElement(paragraph_properties, f"{{{W_NS}}}ind")
        indent.set(f"{{{W_NS}}}firstLine", "708.6614173228347")

    return _rewrite_docx_xml(path, "word/document.xml", _mutate)


def build_with_tables(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Table anchor paragraph.")
    first = doc.add_table(rows=2, cols=2)
    first.style = "Table Grid"
    first.cell(0, 0).text = "Name"
    first.cell(0, 1).text = "Value"
    first.cell(1, 0).text = "Budget"
    first.cell(1, 1).text = "100"
    second = doc.add_table(rows=2, cols=2)
    second.cell(0, 0).text = "Region"
    second.cell(0, 1).text = "Status"
    second.cell(1, 0).text = "West"
    second.cell(1, 1).text = "Ready"
    return _save(doc, path)


def build_table_before_target(path: Path) -> Path:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    table.cell(1, 0).text = "Cell C"
    table.cell(1, 1).text = "Cell D"
    doc.add_paragraph("Canonical anchor paragraph.")
    doc.add_paragraph("UNIQUE_TARGET_AFTER_TABLE")
    return _save(doc, path)


async def _apply_review_operations(path: Path) -> None:
    server = create_server()
    operations = [
        (
            "replace_text",
            {
                "path": str(path),
                "find_text": "draft",
                "replace_with": "approved",
                "track_changes": True,
                "author": "QA",
            },
        ),
        (
            "add_comment_to_matching_text",
            {
                "path": str(path),
                "target_text": "budget",
                "comment_text": "Confirm the source for this budget.",
                "author": "Editor",
                "initials": "ED",
            },
        ),
        (
            "insert_paragraph",
            {
                "path": str(path),
                "after_paragraph": 1,
                "text": "Inserted review note.",
                "track_changes": True,
                "author": "QA",
            },
        ),
    ]
    for name, arguments in operations:
        _, result = await server.call_tool(name, arguments)
        if result["status"] != "ok":
            error_code = result.get("error", {}).get("code", "unknown")
            detail = result.get("error", {}).get("message", result)
            raise RuntimeError(f"Failed to build review fixture with {name} (code: {error_code}): {detail}")


def build_review(path: Path) -> Path:
    build_simple(path)
    asyncio.run(_apply_review_operations(path))
    return path


def build_complex(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Complex sample", level=1)
    doc.add_paragraph("Alpha project budget is 100.")
    doc.add_paragraph("The draft conclusion needs review.")
    callout = doc.styles.add_style("ComplexCallout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.bold = True
    styled = doc.add_paragraph("Styled paragraph in complex sample.")
    styled.style = "ComplexCallout"
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Risk"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Budget"
    table.cell(1, 1).text = "Medium"
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("Second section paragraph.")
    return _save(doc, path)


def generate_all_fixtures(target_dir: Path) -> None:
    builders = {
        "simple.docx": build_simple,
        "with_styles.docx": build_with_styles,
        "with_sections.docx": build_with_sections,
        "fractional_sections.docx": build_fractional_sections,
        "fractional_paragraph_formatting.docx": build_fractional_paragraph_formatting,
        "with_tables.docx": build_with_tables,
        "table_before_target.docx": build_table_before_target,
        "with_review.docx": build_review,
        "complex.docx": build_complex,
    }
    target_dir.mkdir(parents=True, exist_ok=True)
    backup_dir = target_dir / ".extended-docx-mcp-backups"
    if all((target_dir / filename).exists() for filename in builders) and not backup_dir.exists():
        return
    for filename, builder in builders.items():
        builder(target_dir / filename)
    if backup_dir.exists():
        shutil.rmtree(backup_dir)
